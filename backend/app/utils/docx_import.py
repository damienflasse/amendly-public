"""
DOCX import utility — converts a .docx file to clean HTML for TipTap.

Mapping:
  Heading 1              → <h2>  (TipTap top-level heading)
  Heading 2 / 3          → <h3>
  List Bullet            → <ul><li>
  List Number            → <ol><li>
  Normal / other         → <p>

Inline formatting:
  bold                   → <strong>
  italic                 → <em>
  bold + italic          → <strong><em>…</em></strong>
  hyperlinks             → <a href="…">…</a>

Only paragraphs with text content are emitted (blank paragraphs are skipped).
Tables are not imported; their presence is signalled via the warnings field.

Functions:
    docx_bytes_to_html — convert raw .docx bytes to an HTML string.
    DocxImportResult   — dataclass holding html and warnings.
"""

from __future__ import annotations

import html as _html
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

MAX_BYTES = 10 * 1024 * 1024  # 10 MB hard limit


@dataclass
class DocxImportResult:
    """Structured result returned by the DOCX import pipeline."""

    html: str
    warnings: list[str] = field(default_factory=list)


def docx_bytes_to_html(file_bytes: bytes) -> DocxImportResult:
    """
    Convert raw .docx bytes to a clean HTML string with optional warnings.

    The resulting HTML contains a safe subset of tags: p, h2, h3, strong,
    em, ul, ol, li, a.  Tables are skipped; their presence is noted in
    the warnings field.

    Parameters:
        file_bytes: Raw bytes of a .docx file (not base64-encoded).

    Returns:
        DocxImportResult with html string and warnings list.

    Raises:
        ValueError: If the file is too large or cannot be parsed as a .docx.
    """
    if len(file_bytes) > MAX_BYTES:
        raise ValueError("Fichier trop volumineux. Taille maximale : 10 Mo.")

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(
            "Could not parse the uploaded file as a Word document (.docx)."
        ) from exc

    warnings: list[str] = []

    # Detect tables — they are not imported but we warn the user.
    if doc.tables:
        warnings.append("tables_ignored")

    html_parts: list[str] = []
    current_list: str | None = None  # "ul" | "ol" | None

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text_content = para.text  # raw, for emptiness check
        inner_html = _runs_to_html(para)

        # Detect list type from style name or numPr XML element
        list_type = _detect_list_type(para, style_name)

        # ── Close open list when exiting a list block ──
        if not list_type and current_list:
            html_parts.append(f"</{current_list}>")
            current_list = None

        # ── Skip entirely empty non-list paragraphs ──
        if not text_content.strip() and not list_type:
            continue

        if list_type:
            # Open a new list tag if the type changed
            if list_type != current_list:
                if current_list:
                    html_parts.append(f"</{current_list}>")
                html_parts.append(f"<{list_type}>")
                current_list = list_type
            indent_level = _detect_indent_level(para)
            level_attr = f' data-level="{indent_level}"' if indent_level > 0 else ""
            html_parts.append(f"<li{level_attr}>{inner_html}</li>")

        elif "Heading 1" in style_name:
            html_parts.append(f"<h2>{inner_html}</h2>")

        elif "Heading 2" in style_name or "Heading 3" in style_name:
            html_parts.append(f"<h3>{inner_html}</h3>")

        elif inner_html:
            html_parts.append(f"<p>{inner_html}</p>")

    # Close any open list at end of document
    if current_list:
        html_parts.append(f"</{current_list}>")

    return DocxImportResult(html="\n".join(html_parts), warnings=warnings)


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------


def _runs_to_html(para) -> str:
    """
    Convert a paragraph's runs and hyperlinks to an inline HTML string.

    Iterates the paragraph's XML children in document order to preserve
    the interleaving of plain runs and hyperlinks.

    Parameters:
        para: docx Paragraph object.

    Returns:
        HTML string with bold/italic markup and <a href> hyperlinks.
    """
    parts: list[str] = []

    for child in para._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "hyperlink":
            # Extract the relationship URL from the document part
            rId = child.get(qn("r:id"))
            url = ""
            if rId and hasattr(para.part, "rels") and rId in para.part.rels:
                rel = para.part.rels[rId]
                url = rel.target_ref or ""
            # Only allow http(s) and mailto schemes
            if url and (url.startswith("http://") or url.startswith("https://") or url.startswith("mailto:")):
                link_text = "".join(
                    r.text for r in child.iter(qn("w:t")) if r.text
                )
                if link_text:
                    parts.append(f'<a href="{_html.escape(url, quote=True)}">{_html.escape(link_text)}</a>')
            else:
                # Unsupported URL scheme: emit plain text
                link_text = "".join(
                    r.text for r in child.iter(qn("w:t")) if r.text
                )
                if link_text:
                    parts.append(_html.escape(link_text))

        elif tag == "r":
            # Plain run — build a temporary Run-like object by reading XML
            texts = [node.text for node in child.iter(qn("w:t")) if node.text]
            text = _html.escape("".join(texts))
            if not text:
                continue
            rPr = child.find(qn("w:rPr"))
            bold = rPr is not None and (
                rPr.find(qn("w:b")) is not None or rPr.find(qn("w:bCs")) is not None
            )
            italic = rPr is not None and (
                rPr.find(qn("w:i")) is not None or rPr.find(qn("w:iCs")) is not None
            )
            if bold and italic:
                text = f"<strong><em>{text}</em></strong>"
            elif bold:
                text = f"<strong>{text}</strong>"
            elif italic:
                text = f"<em>{text}</em>"
            parts.append(text)

    return "".join(parts)


def _detect_list_type(para, style_name: str) -> str | None:
    """
    Determine whether a paragraph is a list item and, if so, its type.

    Priority:
      1. Style name contains "Number" → ordered list ("ol")
      2. Style name contains "Bullet" or "List" → unordered list ("ul")
      3. XML numPr element present (w:numPr inside w:pPr) → unordered ("ul")
         unless the style is numbered.

    Parameters:
        para: docx Paragraph object.
        style_name: Paragraph's style name string.

    Returns:
        "ul", "ol", or None.
    """
    if "Number" in style_name:
        return "ol"
    if "Bullet" in style_name or "List" in style_name:
        return "ul"
    # Fall back to checking the raw XML for a numPr element
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return "ul"
    return None


def _detect_indent_level(para) -> int:
    """
    Return a zero-based indentation level for a list paragraph.

    Uses the w:ilvl element from the list numbering properties when present,
    falling back to left indent bucketing (every 360 twips ≈ one level).
    Returns 0 for no additional nesting.

    Parameters:
        para: docx Paragraph object.

    Returns:
        Integer indentation level (0 = top level, 1 = first nested, …).
    """
    # Prefer the explicit list level from w:numPr/w:ilvl
    try:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                ilvl = numPr.find(qn("w:ilvl"))
                if ilvl is not None:
                    val = ilvl.get(qn("w:val"))
                    if val is not None:
                        return max(0, int(val))
    except Exception:
        pass

    # Fallback: left indent in twips (1 level ≈ 360 twips)
    try:
        left = para.paragraph_format.left_indent
        if left and left > 0:
            return min(int(left) // 360, 5)
    except Exception:
        pass

    return 0
