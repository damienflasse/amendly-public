"""
Export utilities — generate downloadable versions of a consolidated document.

Each function accepts a document title, body string, and an optional list of amendment
dicts to append as a numbered appendix.  The bytes can be streamed directly from the API.

Supported formats:
  - DOCX  — produced with python-docx
  - PDF   — produced with fpdf2
  - TXT   — UTF-8 plain text (no external library needed)
  - CSV   — UTF-8 comma-separated export for spreadsheet workflows
  - JSON  — UTF-8 structured export for integrations / automation

Amendment dict shape (each item in the optional `amendments` list):
    {
        "number":        int,            # 1-based index
        "author":        str,            # display name or email
        "section":       str | None,
        "original_text": str | None,
        "proposed_text": str | None,
        "justification": str | None,
        "status":        str,            # "accepted" | "pending" | ...
        "created_at":    str,            # pre-formatted date string
    }
"""
import csv
import html
import io
import json
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser

from docx import Document as DocxDocument
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos


@dataclass
class _InlineSegment:
    text: str
    bold: bool = False
    italic: bool = False
    href: str | None = None


@dataclass
class _Block:
    kind: str
    segments: list[_InlineSegment] = field(default_factory=list)
    level: int = 0
    ordered: bool = False
    index: int = 0
    indent_level: int = 0


class _RichTextHTMLParser(HTMLParser):
    """Parse the limited rich-text HTML used by the app into renderable blocks."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[_Block] = []
        self._current: _Block | None = None
        self._style_stack: list[dict[str, object]] = [{"bold": False, "italic": False, "href": None}]
        self._list_stack: list[dict[str, int | bool]] = []
        self._blockquote_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[override]
        attrs_map = dict(attrs)
        if tag in {"strong", "b"}:
            self._push_style(bold=True)
            return
        if tag in {"em", "i"}:
            self._push_style(italic=True)
            return
        if tag == "a":
            self._push_style(href=attrs_map.get("href"))
            return
        if tag == "br":
            self._ensure_block()
            self._append_text("\n")
            return
        if tag == "hr":
            self._flush_block()
            self.blocks.append(_Block(kind="hr"))
            return
        if tag == "blockquote":
            self._flush_block()
            self._blockquote_depth += 1
            return
        if tag in {"ul", "ol"}:
            self._flush_block()
            self._list_stack.append({"ordered": tag == "ol", "index": 0})
            return
        if tag == "li":
            self._flush_block()
            if self._list_stack:
                self._list_stack[-1]["index"] = int(self._list_stack[-1]["index"]) + 1
                ordered = bool(self._list_stack[-1]["ordered"])
                index = int(self._list_stack[-1]["index"])
            else:
                ordered = False
                index = 0
            self._current = _Block(
                kind="list_item",
                ordered=ordered,
                index=index,
                indent_level=max(len(self._list_stack) - 1, 0),
            )
            return
        if tag in {"p", "div"}:
            self._flush_block()
            kind = "blockquote" if self._blockquote_depth else "paragraph"
            self._current = _Block(kind=kind)
            return
        if tag in {"h2", "h3", "h4"}:
            self._flush_block()
            self._current = _Block(kind="heading", level=int(tag[1]))

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        if tag in {"strong", "b", "em", "i", "a"}:
            if len(self._style_stack) > 1:
                self._style_stack.pop()
            return
        if tag in {"p", "div", "li", "h2", "h3", "h4"}:
            self._flush_block()
            return
        if tag == "blockquote":
            self._flush_block()
            self._blockquote_depth = max(self._blockquote_depth - 1, 0)
            return
        if tag in {"ul", "ol"}:
            self._flush_block()
            if self._list_stack:
                self._list_stack.pop()

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        if not data:
            return
        if not data.strip() and self._current is None:
            return
        self._ensure_block()
        self._append_text(data)

    def close(self) -> None:
        super().close()
        self._flush_block()

    def _push_style(self, **updates: object) -> None:
        current = dict(self._style_stack[-1])
        current.update(updates)
        self._style_stack.append(current)

    def _ensure_block(self) -> None:
        if self._current is not None:
            return
        kind = "blockquote" if self._blockquote_depth else "paragraph"
        self._current = _Block(kind=kind)

    def _append_text(self, text: str) -> None:
        if self._current is None:
            return
        style = self._style_stack[-1]
        href = style.get("href")
        segment = _InlineSegment(
            text=text,
            bold=bool(style["bold"]),
            italic=bool(style["italic"]),
            href=str(href) if href else None,
        )
        if self._current.segments and self._can_merge(self._current.segments[-1], segment):
            self._current.segments[-1].text += segment.text
        else:
            self._current.segments.append(segment)

    def _flush_block(self) -> None:
        if self._current is None:
            return
        if _block_has_content(self._current):
            self.blocks.append(self._current)
        self._current = None

    @staticmethod
    def _can_merge(left: _InlineSegment, right: _InlineSegment) -> bool:
        return (
            left.bold == right.bold
            and left.italic == right.italic
            and left.href == right.href
        )


def _block_has_content(block: _Block) -> bool:
    if block.kind == "hr":
        return True
    return any(segment.text for segment in block.segments)


def _looks_like_html(text: str | None) -> bool:
    if not text:
        return False
    return bool(re.search(r"<(?:p|div|br|strong|b|em|i|h2|h3|h4|ul|ol|li|blockquote|hr|a)\b", text))


def _parse_rich_text_blocks(text: str) -> list[_Block]:
    parser = _RichTextHTMLParser()
    parser.feed(text)
    parser.close()
    return parser.blocks


def _plain_text_to_blocks(text: str | None) -> list[_Block]:
    if not text:
        return []
    return [
        _Block(kind="paragraph", segments=[_InlineSegment(text=line.rstrip())])
        for line in text.splitlines()
    ]


def _text_to_blocks(text: str | None) -> list[_Block]:
    if not text:
        return []
    if _looks_like_html(text):
        return _parse_rich_text_blocks(text)
    return _plain_text_to_blocks(text)


def _iter_lines_with_blanks(text: str | None) -> list[str]:
    """Return lines while preserving intentional blank lines."""
    if not text:
        return []
    return text.splitlines()


def _append_docx_paragraphs(
    doc: DocxDocument, text: str | None, *, font_size: int = 11, indent_inches: float = 0.0
) -> None:
    """Append text to a DOCX document while preserving line breaks and spacing."""
    lines = _iter_lines_with_blanks(text)
    if not lines:
        return

    for line in lines:
        para = doc.add_paragraph()
        if indent_inches:
            para.paragraph_format.left_indent = Inches(indent_inches)
        run = para.add_run(line.rstrip())
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"


def _append_docx_segments(para, segments: list[_InlineSegment], *, font_size: int = 11) -> None:
    """Append styled runs to a DOCX paragraph."""
    for segment in segments:
        parts = segment.text.split("\n")
        for idx, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                run.font.size = Pt(font_size)
                run.font.name = "Times New Roman"
                run.bold = segment.bold
                run.italic = segment.italic
                if segment.href:
                    run.underline = True
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            if idx < len(parts) - 1:
                para.add_run().add_break()


def _append_docx_blocks(
    doc: DocxDocument, text: str | None, *, default_font_size: int = 11, indent_inches: float = 0.0
) -> None:
    """Append plain text or supported HTML blocks to DOCX while preserving styling."""
    if not text:
        return
    blocks = _text_to_blocks(text)
    for block in blocks:
        if block.kind == "hr":
            doc.add_paragraph("_" * 40)
            continue

        if block.kind == "heading":
            level = min(max(block.level - 1, 1), 3)
            para = doc.add_heading("", level=level)
            size = {2: 14, 3: 13, 4: 12}.get(block.level, 12)
            _append_docx_segments(para, block.segments, font_size=size)
            continue

        para = doc.add_paragraph()
        if indent_inches:
            para.paragraph_format.left_indent = Inches(indent_inches)
        if block.kind == "blockquote":
            para.paragraph_format.left_indent = Inches(indent_inches + 0.25)
        elif block.kind == "list_item":
            para.paragraph_format.left_indent = Inches(indent_inches + 0.25 + (0.2 * block.indent_level))
            prefix = f"{block.index}. " if block.ordered else "• "
            prefix_run = para.add_run(prefix)
            prefix_run.font.size = Pt(default_font_size)
            prefix_run.font.name = "Times New Roman"

        _append_docx_segments(para, block.segments, font_size=default_font_size)


def _append_docx_labeled_block(
    doc: DocxDocument, label: str, text: str | None, *, font_size: int = 11
) -> None:
    """Render a labeled multi-line block in DOCX with an indented body."""
    if not text:
        return

    label_para = doc.add_paragraph()
    label_run = label_para.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(font_size)
    label_run.font.name = "Times New Roman"
    _append_docx_blocks(doc, text, default_font_size=font_size, indent_inches=0.25)


def _append_pdf_paragraphs(
    pdf: FPDF, text: str | None, *, font_size: int = 11, indent: float = 0.0
) -> None:
    """Append text to a PDF while preserving blank lines and optional indentation."""
    lines = _iter_lines_with_blanks(text)
    if not lines:
        return

    start_x = pdf.l_margin + indent
    width = pdf.w - pdf.r_margin - start_x
    for line in lines:
        pdf.set_x(start_x)
        if line:
            pdf.multi_cell(width, 6, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.ln(6)


def _append_pdf_segments(
    pdf: FPDF,
    segments: list[_InlineSegment],
    *,
    font_size: int = 11,
    line_height: int = 6,
    indent: float = 0.0,
    prefix: str = "",
) -> None:
    """Render styled inline text in PDF using font switches and automatic wrapping."""
    start_x = pdf.l_margin + indent
    pdf.set_x(start_x)

    if prefix:
        pdf.set_font("Helvetica", size=font_size)
        pdf.write(line_height, prefix)

    for segment in segments:
        parts = segment.text.split("\n")
        for idx, part in enumerate(parts):
            if part:
                style = ""
                if segment.bold:
                    style += "B"
                if segment.italic:
                    style += "I"
                pdf.set_font("Helvetica", style=style, size=font_size)
                if segment.href:
                    pdf.set_text_color(0, 102, 204)
                else:
                    pdf.set_text_color(0, 0, 0)
                pdf.write(line_height, part, link=segment.href or "")
            if idx < len(parts) - 1:
                pdf.ln(line_height)
                pdf.set_x(start_x)

    pdf.set_text_color(0, 0, 0)
    pdf.ln(line_height)


def _append_pdf_blocks(
    pdf: FPDF, text: str | None, *, default_font_size: int = 11, indent: float = 0.0
) -> None:
    """Append plain text or supported HTML blocks to PDF while preserving styling."""
    if not text:
        return
    blocks = _text_to_blocks(text)
    for block in blocks:
        if block.kind == "hr":
            y = pdf.get_y() + 2
            pdf.line(pdf.l_margin + indent, y, pdf.w - pdf.r_margin, y)
            pdf.ln(6)
            continue
        if block.kind == "heading":
            size = {2: 14, 3: 13, 4: 12}.get(block.level, 12)
            _append_pdf_segments(
                pdf, block.segments, font_size=size, line_height=8, indent=indent
            )
            pdf.ln(1)
            continue
        block_indent = indent
        prefix = ""
        if block.kind == "blockquote":
            block_indent += 8
        elif block.kind == "list_item":
            block_indent += 6 + (block.indent_level * 6)
            prefix = f"{block.index}. " if block.ordered else "- "
        _append_pdf_segments(
            pdf,
            block.segments,
            font_size=default_font_size,
            line_height=6,
            indent=block_indent,
            prefix=prefix,
        )


def _append_pdf_labeled_block(
    pdf: FPDF, label: str, text: str | None, safe_text
) -> None:
    """Render a labeled multi-line block in the PDF export."""
    if not text:
        return

    pdf.set_font("Helvetica", style="B", size=11)
    pdf.multi_cell(0, 6, label, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    _append_pdf_blocks(pdf, safe_text(text), default_font_size=11, indent=8)


def _segments_to_txt(segments: list[_InlineSegment]) -> str:
    """Convert styled segments to readable plain text."""
    parts: list[str] = []
    for segment in segments:
        text = segment.text
        if segment.href:
            stripped = text.strip()
            if stripped and stripped != segment.href:
                text = f"{text} ({segment.href})"
        parts.append(text)
    return "".join(parts)


def _append_txt_blocks(parts: list[str], text: str | None, *, indent: str = "") -> None:
    """Render plain text or supported HTML as readable TXT."""
    if not text:
        return
    for block in _text_to_blocks(text):
        if block.kind == "hr":
            parts.append(f"{indent}{'-' * 40}\n")
            continue
        if block.kind == "heading":
            heading = _segments_to_txt(block.segments).strip()
            parts.append(f"{indent}{heading}\n")
            parts.append(f"{indent}{'-' * len(heading)}\n")
            continue
        if block.kind == "list_item":
            prefix = f"{block.index}. " if block.ordered else "- "
            content = _segments_to_txt(block.segments).replace("\n", f"\n{indent}  ")
            parts.append(f"{indent}{prefix}{content}\n")
            continue
        content = _segments_to_txt(block.segments).replace("\n", f"\n{indent}")
        parts.append(f"{indent}{content}\n")


def _append_txt_labeled_block(parts: list[str], label: str, text: str | None) -> None:
    """Render a labeled multi-line block in plain text with indentation."""
    if not text:
        return

    parts.append(f"{label}\n")
    _append_txt_blocks(parts, text, indent="  ")


def export_docx(title: str, body: str, amendments: list[dict] | None = None) -> bytes:
    """
    Generate a DOCX file from a document title, body, and optional amendments appendix.

    The title is set as a Level-1 heading.  The body is split on newlines and each
    non-empty line is written as a separate paragraph with 11pt Times New Roman body
    text.  If `amendments` is provided and non-empty, a page break is inserted followed
    by an "Amendments" Level-2 heading and a numbered entry for each amendment.

    Parameters:
        title:       Document title (used as the first heading).
        body:        Full consolidated body text (may contain newlines).
        amendments:  Optional list of amendment dicts to append.  Each dict must
                     contain at least 'number', 'author', 'created_at', and one of
                     'original_text'/'proposed_text' or 'justification'.

    Returns:
        Raw bytes of a valid .docx file.
    """
    doc = DocxDocument()

    # Title heading
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

    # Body paragraphs — preserve supported rich-text styling when present
    _append_docx_blocks(doc, body, default_font_size=11)

    # Amendments appendix
    if amendments:
        doc.add_page_break()
        appendix_heading = doc.add_heading("Amendments", level=2)
        for run in appendix_heading.runs:
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"

        for amend in amendments:
            # Sub-heading per amendment: "#N — Section (status)"
            entry_label = f"#{amend['number']}"
            if amend.get("section"):
                entry_label += f" — {amend['section']}"
            entry_label += f"  [{amend['status'].upper()}]"
            entry_h = doc.add_heading(entry_label, level=3)
            for run in entry_h.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"

            # Author / date meta line
            meta = f"By {amend['author']}  ·  {amend['created_at']}"
            meta_para = doc.add_paragraph(meta)
            for run in meta_para.runs:
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

            if amend.get("original_text") is not None:
                _append_docx_labeled_block(doc, "Original:", amend["original_text"], font_size=11)
                _append_docx_labeled_block(doc, "Proposed:", amend["proposed_text"], font_size=11)

            if amend.get("justification"):
                _append_docx_labeled_block(
                    doc, "Justification:", amend["justification"], font_size=11
                )

            # Spacer between amendments
            doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(title: str, body: str, amendments: list[dict] | None = None) -> bytes:
    """
    Generate a PDF file from a document title, body, and optional amendments appendix.

    The title is rendered in bold at 16pt.  Body text is rendered at 11pt with
    multi-cell word-wrap.  If `amendments` is provided and non-empty, a separator is
    added followed by an "Amendments" section with one numbered entry per amendment.
    UTF-8 characters are supported via the built-in helvetica font (latin-1 subset)
    with graceful replacement of characters outside the latin-1 range.

    Parameters:
        title:       Document title.
        body:        Full consolidated body text.
        amendments:  Optional list of amendment dicts to append.

    Returns:
        Raw bytes of a valid .pdf file.
    """
    def _safe(text: str) -> str:
        """Encode to Windows-1252, preserving common editorial punctuation."""
        return text.encode("cp1252", errors="replace").decode("cp1252")

    pdf = FPDF()
    pdf.core_fonts_encoding = "cp1252"
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(left=20, top=20, right=20)

    # Title
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.multi_cell(0, 10, _safe(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Body
    if body:
        _append_pdf_blocks(pdf, _safe(body), default_font_size=11)

    # Amendments appendix
    if amendments:
        pdf.add_page()
        pdf.set_font("Helvetica", style="B", size=14)
        pdf.multi_cell(0, 10, "Amendments", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        for amend in amendments:
            # Entry heading
            entry_label = f"#{amend['number']}"
            if amend.get("section"):
                entry_label += f" - {amend['section']}"
            entry_label += f"  [{amend['status'].upper()}]"

            pdf.set_font("Helvetica", style="B", size=12)
            pdf.multi_cell(0, 8, _safe(entry_label), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Meta line
            meta = f"By {amend['author']}  -  {amend['created_at']}"
            pdf.set_font("Helvetica", style="I", size=10)
            pdf.multi_cell(0, 6, _safe(meta), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)

            if amend.get("original_text") is not None:
                _append_pdf_labeled_block(pdf, "Original:", amend["original_text"], _safe)
                _append_pdf_labeled_block(pdf, "Proposed:", amend["proposed_text"], _safe)

            if amend.get("justification"):
                _append_pdf_labeled_block(
                    pdf, "Justification:", amend["justification"], _safe
                )

            pdf.ln(4)

    return bytes(pdf.output())


def export_txt(title: str, body: str, amendments: list[dict] | None = None) -> bytes:
    """
    Generate a plain-text version of a document with an optional amendments appendix.

    The title is written on the first line, followed by a separator, then the body.
    If `amendments` is provided and non-empty, a section separator and a numbered
    entry per amendment are appended.  The result is UTF-8 encoded.

    Parameters:
        title:       Document title.
        body:        Full consolidated body text.
        amendments:  Optional list of amendment dicts to append.

    Returns:
        UTF-8 encoded bytes of the plain-text document.
    """
    separator = "=" * len(title)
    parts = [f"{title}\n{separator}\n\n"]
    if body:
        _append_txt_blocks(parts, body)

    if amendments:
        if body:
            parts.append("\n")
        parts.append("\n" + "=" * 60 + "\nAMENDMENTS\n" + "=" * 60 + "\n")
        for amend in amendments:
            entry_label = f"\n#{amend['number']}"
            if amend.get("section"):
                entry_label += f" — {amend['section']}"
            entry_label += f"  [{amend['status'].upper()}]"
            parts.append(entry_label + "\n")
            parts.append(f"By {amend['author']}  ·  {amend['created_at']}\n")

            if amend.get("original_text") is not None:
                _append_txt_labeled_block(parts, "Original:", amend["original_text"] or "")
                _append_txt_labeled_block(parts, "Proposed:", amend["proposed_text"] or "")

            if amend.get("justification"):
                _append_txt_labeled_block(parts, "Justification:", amend["justification"])

    return "".join(parts).encode("utf-8")


def export_csv(title: str, body: str, amendments: list[dict] | None = None) -> bytes:
    """
    Generate a CSV export of the document and its optional amendments appendix.

    The first row always represents the consolidated document itself
    (`row_type=document`). Each amendment is then appended as its own
    `row_type=amendment` entry so spreadsheet users can filter or pivot the
    dataset without needing a second file.

    Parameters:
        title:       Document title.
        body:        Full consolidated body text.
        amendments:  Optional list of amendment dicts to append as rows.

    Returns:
        UTF-8 encoded CSV bytes.
    """
    buffer = io.StringIO(newline="")
    writer = csv.DictWriter(
        buffer,
        fieldnames=[
            "row_type",
            "title",
            "body",
            "number",
            "author",
            "section",
            "original_text",
            "proposed_text",
            "justification",
            "status",
            "created_at",
        ],
    )
    writer.writeheader()
    writer.writerow(
        {
            "row_type": "document",
            "title": title,
            "body": body or "",
            "number": "",
            "author": "",
            "section": "",
            "original_text": "",
            "proposed_text": "",
            "justification": "",
            "status": "",
            "created_at": "",
        }
    )

    for amend in amendments or []:
        writer.writerow(
            {
                "row_type": "amendment",
                "title": title,
                "body": "",
                "number": amend.get("number", ""),
                "author": amend.get("author", ""),
                "section": amend.get("section", ""),
                "original_text": amend.get("original_text", ""),
                "proposed_text": amend.get("proposed_text", ""),
                "justification": amend.get("justification", ""),
                "status": amend.get("status", ""),
                "created_at": amend.get("created_at", ""),
            }
        )

    return buffer.getvalue().encode("utf-8")


def export_json(title: str, body: str, amendments: list[dict] | None = None) -> bytes:
    """
    Generate a JSON export of the consolidated document and optional amendments.

    Parameters:
        title:       Document title.
        body:        Full consolidated body text.
        amendments:  Optional list of amendment dicts to append.

    Returns:
        UTF-8 encoded JSON bytes.
    """
    payload = {
        "title": title,
        "body": body or "",
        "amendments": amendments or [],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
